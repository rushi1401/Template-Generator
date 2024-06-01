import streamlit as st
from langchain_groq import ChatGroq
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFDirectoryLoader
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.chains import ConversationalRetrievalChain
from docx import Document

# Initialize the Groq LLM
groq_api_key = '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'
llm = ChatGroq(groq_api_key=groq_api_key, model_name="Llama3-8b-8192")

# Embedding using Huggingface
huggingface_embeddings = HuggingFaceEmbeddings(
    model_name="BAAI/bge-small-en-v1.5",
    model_kwargs={'device': 'cpu'},
    encode_kwargs={'normalize_embeddings': True}
)

# Function to load FAISS vector store based on selected sector
def load_faiss_database(sector):
    if sector == "Health":
        return FAISS.load_local("faiss_index_medical", huggingface_embeddings, allow_dangerous_deserialization=True)
    elif sector == "Banking":
        return FAISS.load_local("faiss_index_banking", huggingface_embeddings, allow_dangerous_deserialization=True)
    elif sector == "Real Estate":
        return FAISS.load_local("faiss_index_realestate", huggingface_embeddings, allow_dangerous_deserialization=True)

# Function to fetch answers for multiple keywords and save to DOCX
def get_answers_and_save_docx(chain, keywords):
    doc = Document()
    doc.add_heading('Templete Example', 0)
    
    chat_history = []
    for keyword in keywords:
        result = chain({"question": keyword, "chat_history": chat_history})
        answer = result['answer']
        
        doc.add_heading(f'Keyword: {keyword}', level=1)
        doc.add_paragraph(answer)
        
    doc.save("Templete_Gen.docx")

# Function to read keywords from a template file
def read_keywords_from_template(template_path):
    with open(template_path, 'r') as file:
        keywords = file.readlines()
    return [keyword.strip() for keyword in keywords if keyword.strip()]

# Streamlit app
def main():
    st.title("Templete Generator")

    # Sidebar for sector selection
    sector = st.sidebar.selectbox("Select Sector", ["Health", "Banking", "Real Estate"])

    # Path to the template file
    template_path = 'template.txt'

    # Read keywords from the template file
    keywords = read_keywords_from_template(template_path)

    if st.button("Get Templete"):
        if keywords:
            # Load the appropriate FAISS vector store based on selected sector
            new_db = load_faiss_database(sector)
            chain = ConversationalRetrievalChain.from_llm(llm, new_db.as_retriever(), return_source_documents=True)
            
            with st.spinner("Fetching the answers..."):
                get_answers_and_save_docx(chain, keywords)
                
                st.success("Answers saved to Chatbot_Answers.docx")
                with open("Templete_Gen.docx", "rb") as file:
                    btn = st.download_button(
                        label="Download DOCX",
                        data=file,
                        file_name="Templete_Answers.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else:
            st.warning("No keywords found in the template file.")

if __name__ == "__main__":
    main()
